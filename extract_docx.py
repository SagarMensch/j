import docx

doc = docx.Document(r"c:\Users\sagar\Downloads\jubilantingrevia\Discussion on BRD SOP.docx")
md_lines = []

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        md_lines.append("")
        continue

    style = para.style.name if para.style else ""

    if "Heading 1" in style:
        md_lines.append(f"# {text}")
    elif "Heading 2" in style:
        md_lines.append(f"## {text}")
    elif "Heading 3" in style:
        md_lines.append(f"### {text}")
    elif "Heading 4" in style:
        md_lines.append(f"#### {text}")
    elif "List" in style or para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr') is not None:
        md_lines.append(f"- {text}")
    else:
        # Check if entire paragraph is bold
        is_bold = para.runs and all(run.bold for run in para.runs if run.text.strip())
        if is_bold:
            md_lines.append(f"**{text}**")
        else:
            md_lines.append(text)

# Handle tables
for table in doc.tables:
    md_lines.append("")
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        md_lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            md_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
    md_lines.append("")

with open(r"c:\Users\sagar\Downloads\jubilantingrevia\Discussion on BRD SOP.md", "w", encoding="utf-8") as f:
    f.write("\n".join(md_lines))

print("Done! Saved as 'Discussion on BRD SOP.md'")
